"""Apply small layout fixes to the Pandoc Word export."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Twips
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


TABLE_ONE_CAPTION = "Method and controls on one scale."
CONTENT_TABLE_CAPTION = "Cosmos 3 isolated-content interventions"
PATHWAY_TABLE_CAPTION = "Cosmos 3 future-token pathway intervention"
PARAGRAPH_LABELS = {
    "Cosmos Policy.",
    "Cosmos 3.",
    "Limitations.",
    "Broader impact.",
}


def keep_row_together(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        properties.append(header)


def set_table_geometry(table, widths: tuple[int, ...]) -> None:
    """Keep Word and LibreOffice on the same fixed-width table geometry."""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr

    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(widths)))

    table_indent = properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), "120")

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for index, width in enumerate(widths):
        table.columns[index].width = Twips(width)
    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths)):
            cell.width = Twips(width)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(width))


def main(path: Path) -> None:
    document = Document(path)
    for paragraph in document.paragraphs:
        label = re.sub(r"^\d+(?:\.\d+)*\t", "", paragraph.text.strip())
        if label not in PARAGRAPH_LABELS:
            continue
        paragraph.clear()
        paragraph.add_run(label)
        paragraph.style = document.styles["Normal"]
        properties = paragraph._p.get_or_add_pPr()
        numbering = properties.find(qn("w:numPr"))
        if numbering is not None:
            properties.remove(numbering)
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            run.font.bold = True
            run.font.italic = False
            run.font.color.rgb = RGBColor(31, 99, 125)

    caption = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(TABLE_ONE_CAPTION)
    )
    caption.paragraph_format.page_break_before = True
    caption.paragraph_format.keep_with_next = True

    content_caption = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(CONTENT_TABLE_CAPTION)
    )
    content_caption.paragraph_format.page_break_before = True
    content_caption.paragraph_format.keep_with_next = True

    pathway_caption = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(PATHWAY_TABLE_CAPTION)
    )
    pathway_caption.paragraph_format.keep_with_next = True

    primary_table = document.tables[0]
    set_table_geometry(primary_table, (1350, 950, 850, 1700, 4510))
    for row in primary_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    repeat_header(primary_table.rows[0])
    for row in primary_table.rows:
        keep_row_together(row)

    content_table = document.tables[1]
    set_table_geometry(content_table, (2160, 3600, 3600))
    for row in content_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    repeat_header(content_table.rows[0])
    for row in content_table.rows:
        keep_row_together(row)

    pathway_table = document.tables[2]
    set_table_geometry(pathway_table, (1944, 1728, 3816, 1872))
    for row in pathway_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    repeat_header(pathway_table.rows[0])
    for row in pathway_table.rows:
        keep_row_together(row)

    document.save(path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("usage: postprocess_docx.py PATH")
    main(Path(sys.argv[1]))
